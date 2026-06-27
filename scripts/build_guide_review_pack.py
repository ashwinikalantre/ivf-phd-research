from pathlib import Path
import csv
import io

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "guide_review_pack"


def load_literature_rows():
    matrix = ROOT / "docs" / "evidence-base" / "literature-matrix.md"
    rows = []
    for line in matrix.read_text(encoding="utf-8").splitlines():
        if not line.startswith("| "):
            continue
        if line.startswith("| No ") or line.startswith("| ---"):
            continue
        raw = line.strip().strip("|")
        parsed = next(csv.reader(io.StringIO(raw), delimiter="|", skipinitialspace=True))
        cells = [cell.strip() for cell in parsed]
        if len(cells) != 9:
            continue
        rows.append(
            {
                "no": cells[0],
                "year": cells[1],
                "title": cells[2],
                "journal": cells[3],
                "objective": cells[4],
                "method": cells[5],
                "dataset": cells[6],
                "gap": cells[7],
                "relevance": cells[8],
            }
        )
    return rows


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, font_size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(font_size)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        set_cell_shading(hdr[i], "F2F4F7")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_compact_paper_table(doc, rows):
    headers = ["Paper", "Focus and Method", "Dataset / Evidence", "Main Gap / Use"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, font_size=8)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        values = [
            f'Paper {row["no"]} ({row["year"]})\n{row["title"]}\n{row["journal"]}',
            f'{row["objective"]}\nMethod: {row["method"]}',
            row["dataset"],
            f'{row["gap"]}\nRelevance: {row["relevance"]}',
        ]
        for i, value in enumerate(values):
            set_cell_text(cells[i], value, font_size=8)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def setup_doc(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(11, 37, 69)

    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(subtitle)
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_paragraph()
    return doc


def save_doc(doc, filename):
    path = OUT / filename
    doc.save(path)
    return path


def literature_summary():
    doc = setup_doc(
        "Compact Literature Review Summary",
        "Guide review pack | AI, ML, XAI and decision support in IVF | Working evidence base",
    )
    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "This document summarizes the current literature-review work for guide discussion. "
        "It is not a final thesis chapter. Detailed extraction remains in the Excel workbook and individual paper notes."
    )
    doc.add_heading("Review Scope", level=1)
    add_table(
        doc,
        ["Item", "Current Status"],
        [
            ["Domain", "Women’s healthcare, focused on IVF and assisted reproductive technology."],
            ["Technical area", "AI, machine learning, explainable AI, clinical decision support and personalized prediction."],
            ["Primary time window", "2021-2025, with a few highly relevant 2026 synthesis/model papers kept as supporting evidence."],
            ["Current reviewed notes", "35 structured paper notes."],
            ["Main source of detail", "Excel literature review matrix and paper-wise documentation pages."],
        ],
        [1.7, 4.6],
    )
    doc.add_heading("Strong Literature Signals", level=1)
    add_bullets(
        doc,
        [
            "Many recent IVF-AI papers focus on prediction of live birth, clinical pregnancy, implantation, miscarriage or ovarian response.",
            "Several studies use tabular clinical data with models such as logistic regression, random forest, XGBoost and LightGBM.",
            "Embryo-image and embryo-selection AI are active areas, but stronger datasets and clinical validation are often required.",
            "Explainability is repeatedly discussed as important for clinical trust, patient communication and ethical AI use.",
            "Decision-support framing is safer than claiming that AI can replace clinical judgment.",
        ],
    )
    doc.add_heading("Representative Paper Groups", level=1)
    add_table(
        doc,
        ["Group", "Examples", "Use in Topic Discovery"],
        [
            ["Live-birth / pregnancy prediction", "Papers 5, 6, 7, 8, 9, 29, 30, 31, 33", "Supports outcome-prediction direction."],
            ["XAI / trustworthy AI", "Papers 1, 14, 25, 26, 27, 33", "Supports explainable model and explanation-output direction."],
            ["CDSS / personalization", "Papers 2, 21, 22, 23", "Supports clinical decision-support and treatment-personalization direction."],
            ["Embryo AI", "Papers 3, 4, 24, 26, 27, 28", "Useful if embryo variables or images are available."],
            ["Review / trend papers", "Papers 12, 13, 15, 32, 34, 35", "Useful for background, gaps and justification."],
        ],
        [1.45, 1.9, 2.95],
    )
    doc.add_heading("Current Interpretation", level=1)
    doc.add_paragraph(
        "The strongest working direction is explainable and personalized IVF outcome prediction with a clinical decision-support framing. "
        "The final title should remain open until clinic data feasibility is confirmed."
    )
    doc.add_page_break()
    doc.add_heading("Paper-Wise Compact Summary", level=1)
    doc.add_paragraph(
        "This section gives a compact guide-facing summary of the current 35 reviewed papers. "
        "It is derived from the working literature matrix; full extraction fields remain in the Excel workbook and paper-wise notes."
    )
    add_compact_paper_table(doc, load_literature_rows())
    return save_doc(doc, "02_compact_literature_review_summary.docx")


def theme_gap_analysis():
    doc = setup_doc(
        "Theme and Gap Analysis",
        "Guide review pack | How the current topic direction emerged from repeated literature signals",
    )
    doc.add_heading("Why Themes Were Created", level=1)
    doc.add_paragraph(
        "Theme classification helps convert individual papers into research directions. "
        "It prevents title selection from depending on a single paper or isolated limitation."
    )
    doc.add_heading("Theme Classification Logic", level=1)
    doc.add_paragraph(
        "The review was organized as: paper-wise extraction, theme grouping, repeated-gap analysis, research opportunity identification, "
        "and then candidate topic development. The final title should emerge from repeated gaps and dataset feasibility, not from a single attractive keyword."
    )
    doc.add_heading("Major Theme Groups", level=1)
    add_table(
        doc,
        ["Major Theme", "Paper Examples", "How It Helps Topic Selection", "Dataset Dependency"],
        [
            [
                "IVF outcome prediction",
                "5, 6, 7, 8, 9, 15, 29, 30, 31, 33",
                "Provides the core predictive analytics direction: clinical pregnancy, live birth, miscarriage or related IVF outcome.",
                "Needs a clear outcome variable; live birth is strongest but may be harder to collect.",
            ],
            [
                "Explainable / trustworthy AI",
                "1, 14, 25, 26, 27, 33",
                "Supports moving beyond black-box prediction toward interpretable, patient-level and clinician-facing explanations.",
                "Needs model outputs that can be explained; clinician review is ideal.",
            ],
            [
                "Clinical decision support",
                "2, 21, 22, 23, 34",
                "Positions the PhD as a usable framework, not only an accuracy comparison study.",
                "Needs careful framing: support clinical counseling or planning, not replace doctors.",
            ],
            [
                "Personalization",
                "1, 2, 12, 21, 23",
                "Supports patient-specific prediction and treatment-context interpretation.",
                "Needs enough patient-level variables and subgroup size.",
            ],
        ],
        [1.35, 1.1, 2.55, 1.35],
    )
    doc.add_heading("Supporting Themes", level=1)
    add_table(
        doc,
        ["Supporting Theme", "Paper Examples", "Use", "Caution"],
        [
            [
                "Embryo and multimodal AI",
                "3, 4, 24, 27, 28",
                "Useful if embryology variables, embryo grades, images or time-lapse data are available.",
                "Do not promise image/deep-learning work without data access.",
            ],
            [
                "Ovarian stimulation / protocol decision support",
                "1, 2, 21, 22, 23",
                "Strong direction if dose, trigger and stimulation-monitoring records are available.",
                "May require more detailed clinical workflow data than basic outcome prediction.",
            ],
            [
                "Indian/context-specific validation",
                "1, 5, 29",
                "Potentially important because many models are not validated in Indian IVF settings.",
                "Should appear in title only if Indian clinic data and validation are feasible.",
            ],
            [
                "Lifestyle and contextual data",
                "Mainly from planned data framework, not yet a strong repeated paper gap",
                "Can enrich personalization if reliably collected.",
                "Should remain conditional until questionnaire/data feasibility is confirmed.",
            ],
        ],
        [1.45, 1.15, 2.25, 1.5],
    )
    doc.add_heading("Repeated Gap Assessment", level=1)
    add_table(
        doc,
        ["Gap", "Frequency", "Evidence Strength", "Actionable in PhD?", "Topic Relevance"],
        [
            [
                "Need external validation",
                "26",
                "Very strong",
                "Conditional",
                "Strong if independent clinic, source or time-period data is available.",
            ],
            [
                "Need clinical decision support",
                "18",
                "Very strong",
                "Yes",
                "Supports CDSS framing and avoids an accuracy-only thesis.",
            ],
            [
                "Need real-world deployment evidence",
                "9",
                "Strong",
                "Partly",
                "Useful for justification; full deployment may be future work.",
            ],
            [
                "Retrospective design",
                "9",
                "Strong",
                "Partly",
                "Can be handled through careful validation and honest limitation wording.",
            ],
            [
                "Trustworthy AI gap",
                "8",
                "Strong",
                "Yes",
                "Supports XAI, clinician review and conservative claims.",
            ],
            [
                "Single-center study",
                "6",
                "Moderate-strong",
                "Conditional",
                "Supports multi-center or temporal validation if feasible.",
            ],
        ],
        [1.4, 0.7, 1.05, 1.0, 2.2],
    )
    doc.add_heading("Strong Gaps Versus Conditional Gaps", level=1)
    add_table(
        doc,
        ["Gap Type", "Gaps", "How To Treat In Proposal"],
        [
            [
                "Strong repeated gaps",
                "External validation, clinical decision support, trustworthy/explainable AI, real-world clinical usefulness, retrospective/single-center limitations.",
                "Use these to justify the main research direction.",
            ],
            [
                "Conditional gaps",
                "Indian validation, lifestyle data, embryo-image/deep-learning work, multi-center validation.",
                "Use only if data access supports them.",
            ],
            [
                "Weak/not-yet-claimable gaps",
                "Any gap supported by only one paper or by our interest alone.",
                "Do not use as novelty claim until stronger evidence is collected.",
            ],
        ],
        [1.35, 2.6, 2.4],
    )
    doc.add_heading("Gaps To Treat Carefully", level=1)
    add_bullets(
        doc,
        [
            "Lifestyle data is promising, but should not be claimed as a repeated core gap unless clinic/questionnaire feasibility is confirmed.",
            "Indian population validation is valuable, but should appear in the final title only if Indian clinic data is available.",
            "Embryo-image deep learning should not be promised unless image or time-lapse data can be accessed ethically and practically.",
        ],
    )
    doc.add_heading("Working Research Opportunity", level=1)
    doc.add_paragraph(
        "The strongest current research opportunity is not simply to build another IVF prediction model. "
        "A more defensible PhD direction is to design an explainable, personalized and clinically usable decision-support framework "
        "for IVF outcome prediction, with the exact data modality and validation strategy finalized after clinic data feasibility is known."
    )
    doc.add_heading("What Should Not Be Claimed Yet", level=1)
    add_bullets(
        doc,
        [
            "Do not claim that AI will improve IVF success rates unless a clinical outcome study proves it.",
            "Do not claim external validation unless an independent source, clinic or time period is used.",
            "Do not include Indian women, lifestyle data or embryo images in the final title unless those data are confirmed.",
            "Do not present model accuracy as the only PhD contribution; explanation, validation and clinical usability matter.",
        ],
    )
    doc.add_heading("Guide Discussion Questions", level=1)
    add_numbered(
        doc,
        [
            "Which theme should be prioritized: outcome prediction, stimulation decision support or broader clinical decision support?",
            "Should the topic remain broad until data access is confirmed, or should it be narrowed now?",
            "Should Indian validation be part of the title, or only a conditional objective?",
            "Which gaps are strong enough for the proposal, and which should remain future work?",
            "What dataset scenario is realistic for the next meeting with doctors or clinics?",
        ],
    )
    return save_doc(doc, "03_theme_and_gap_analysis.docx")


def dataset_feasibility():
    doc = setup_doc(
        "Dataset Feasibility and Data Requirement Plan",
        "Guide review pack | Variables and feasibility questions before final title selection",
    )
    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "The final PhD topic depends heavily on clinic data access. This document converts dataset uncertainty into a decision plan: "
        "what minimum data are needed, what stronger data would improve the study, and which title words should remain conditional."
    )
    doc.add_heading("Core Decision Gates", level=1)
    add_table(
        doc,
        ["Decision Question", "If Yes", "If No"],
        [
            [
                "Can IVF clinical data be accessed?",
                "A basic explainable IVF outcome prediction study is feasible.",
                "The current topic cannot move forward practically; need another clinic/source.",
            ],
            [
                "Is a clear outcome available?",
                "Model target can be defined: clinical pregnancy, live birth, miscarriage or oocyte/embryo response.",
                "Prediction study is weak; outcome definition must be solved first.",
            ],
            [
                "Are embryology variables available?",
                "The title may include multimodal clinical and embryological data.",
                "Use clinical data only; keep embryology as future scope.",
            ],
            [
                "Are lifestyle variables available or collectable?",
                "Lifestyle personalization can be explored carefully.",
                "Do not include lifestyle in the title.",
            ],
            [
                "Is independent validation possible?",
                "External, clinic-wise or temporal validation can be claimed depending on source.",
                "Use internal validation and clearly state limitation.",
            ],
            [
                "Can doctors review explanation outputs?",
                "Clinician usefulness/usability can become an objective.",
                "Keep CDSS as a proposed framework, not clinically evaluated deployment.",
            ],
        ],
        [1.95, 2.2, 2.15],
    )
    doc.add_heading("Variable Priority Levels", level=1)
    add_table(
        doc,
        ["Level", "Meaning", "Example"],
        [
            ["Mandatory", "Minimum variables required for a defensible IVF outcome-prediction model.", "Age, BMI, infertility diagnosis, AMH, AFC, treatment type, outcome."],
            ["Recommended", "Strongly useful for personalization, model quality or explanation.", "Previous IVF attempts, gonadotropin dose, oocytes retrieved, embryo grade."],
            ["Optional", "Useful if available, but not essential for the first study.", "Diet, sleep, stress, ultrasound images, PGT-A status."],
            ["Sensitive", "Potentially valuable but needs stronger privacy, consent and ethics handling.", "Income, occupation, doctor ID, embryologist ID, images, genetic/ploidy data."],
        ],
        [1.2, 2.65, 2.45],
    )
    doc.add_page_break()
    doc.add_heading("Minimum Viable Dataset", level=1)
    add_table(
        doc,
        ["Category", "Minimum Variables", "Why It Matters"],
        [
            ["Outcome", "Clinical pregnancy or live birth.", "At least one clear dependent variable is required."],
            ["Demographic", "Age, BMI, infertility duration, infertility type/cause.", "Core patient-level predictors and counselling variables."],
            ["Clinical/hormonal", "AMH, AFC, diagnosis such as PCOS/endometriosis if available.", "Represents ovarian reserve and major clinical subgroups."],
            ["Treatment", "IVF/ICSI, fresh/frozen transfer, stimulation protocol, endometrial thickness, embryos transferred.", "Captures treatment pathway and transfer context."],
            ["Validation", "Train/test, cross-validation or temporal split.", "External validation only if independent data exists."],
        ],
        [1.25, 2.75, 2.25],
    )
    doc.add_heading("Strong Dataset Scenario", level=1)
    add_table(
        doc,
        ["Data Group", "Useful Variables", "Topic Benefit"],
        [
            ["Embryology", "Oocytes retrieved, mature oocytes, fertilization rate, embryo grade, blastocyst grade, transfer day.", "Allows multimodal clinical-embryological prediction."],
            ["Treatment intensity", "Starting dose, total gonadotropin dose, stimulation duration, trigger type/day.", "Supports personalization and possible stimulation-response analysis."],
            ["Male partner", "Sperm count/concentration, motility, morphology, male-factor diagnosis.", "Improves couple-level IVF context if available."],
            ["Time/workflow", "Treatment year, protocol changes, clinic ID if available.", "Supports temporal or clinic-wise validation and model drift discussion."],
        ],
        [1.3, 3.0, 2.0],
    )
    doc.add_heading("Ideal Dataset Scenario", level=1)
    add_table(
        doc,
        ["Ideal Addition", "Use", "Caution"],
        [
            ["Lifestyle questionnaire", "Smoking, physical activity, sleep, stress, diet, caffeine or occupation.", "Use validated or clearly defined measures; avoid weak self-report claims."],
            ["Multiple clinics or independent time period", "External, clinic-wise or temporal validation.", "Do not call internal testing external validation."],
            ["Clinician review", "Checks whether explanation outputs are meaningful for counselling.", "This evaluates usefulness, not clinical outcome improvement."],
            ["Images/time-lapse/radiomics", "Possible deep learning or radiomics direction.", "Only feasible with storage, consent, annotation and technical support."],
        ],
        [1.7, 2.55, 2.05],
    )
    doc.add_heading("Outcome Selection", level=1)
    add_table(
        doc,
        ["Outcome", "Priority", "Use In Study", "Caution"],
        [
            ["Live birth", "Preferred if available", "Strongest patient-relevant IVF endpoint.", "May require follow-up and complete records."],
            ["Clinical pregnancy", "Minimum acceptable", "Often easier to obtain and still clinically meaningful.", "Less final than live birth."],
            ["Miscarriage", "Secondary", "Useful for post-pregnancy risk modelling.", "Only relevant after pregnancy is confirmed."],
            ["Oocyte/mature oocyte yield", "Secondary", "Useful for ovarian stimulation response.", "Different target from pregnancy/live birth."],
            ["Good-quality embryo/blastocyst", "Secondary", "Useful if embryology data exists.", "Lab outcome, not final patient outcome."],
        ],
        [1.3, 1.25, 2.3, 1.5],
    )
    doc.add_heading("Validation Options", level=1)
    add_table(
        doc,
        ["Validation Type", "When It Is Possible", "How To Word It Safely"],
        [
            ["Internal validation", "One dataset from one clinic.", "Internal model validation only."],
            ["Temporal validation", "Same clinic, later years held out.", "Validation on a later time period."],
            ["External validation", "Different clinic/source/system.", "External validation, if truly independent."],
            ["Subgroup validation", "Adequate sample sizes by age, diagnosis, PCOS/endometriosis etc.", "Subgroup performance analysis, not proof of fairness unless designed properly."],
        ],
        [1.45, 2.35, 2.5],
    )
    doc.add_heading("Ethics and Privacy Requirements", level=1)
    add_bullets(
        doc,
        [
            "Use anonymized or de-identified patient/cycle IDs; direct identifiers should not enter the research dataset.",
            "Confirm whether ethics committee approval, clinic permission and data-use agreement are required.",
            "Handle sensitive data carefully: income, occupation, doctor/embryologist IDs, embryo images, genetic/ploidy data and lifestyle details.",
            "Keep the system framed as clinical decision support; the model should not make autonomous treatment decisions.",
        ],
    )
    doc.add_heading("Doctor / Clinic Questions", level=1)
    add_numbered(
        doc,
        [
            "Which IVF variables are routinely recorded in electronic or paper records?",
            "Is live birth available, or only clinical pregnancy?",
            "Approximately how many IVF cycles are available, and from which years?",
            "Are embryology variables available in structured form?",
            "Are stimulation dose, trigger and follicle-monitoring variables available?",
            "Are lifestyle variables already recorded, or would a questionnaire be needed?",
            "Can data from more than one clinic, doctor, year or time period be accessed?",
            "Can clinicians review whether explanation outputs are meaningful?",
            "What anonymization, permission and ethics process will be required?",
        ],
    )
    doc.add_page_break()
    doc.add_heading("Title Implication", level=1)
    add_table(
        doc,
        ["Data Available", "Safe Topic Direction", "Avoid Claiming"],
        [
            ["Clinical data only", "Explainable clinical IVF outcome prediction.", "Multimodal or embryology-based title."],
            ["Clinical + embryology", "Explainable multimodal IVF outcome prediction.", "Image/deep-learning claims unless images are available."],
            ["Clinical + embryology + lifestyle", "Personalized IVF prediction using clinical, embryological and lifestyle data.", "Lifestyle effect claims without reliable measurement."],
            ["Independent clinic/source", "External validation in Indian IVF settings can be considered.", "External validation if only one-center internal split exists."],
            ["Clinician review possible", "Clinician-facing explanation/usability evaluation can be added.", "Clinical deployment or improved success rate."],
        ],
        [1.8, 2.55, 1.95],
    )
    doc.add_heading("Current Safe Position For Guide", level=1)
    doc.add_paragraph(
        "The safest current position is that the PhD direction is dataset-dependent. "
        "A clinical-data-only study is still possible, but the stronger topic is clinical plus embryological data with explainability and decision-support framing. "
        "Indian validation, lifestyle data, external validation and clinician usability should remain conditional until data access is confirmed."
    )
    return save_doc(doc, "04_dataset_feasibility_and_data_requirements.docx")


def methodology_blueprint():
    doc = setup_doc(
        "Methods, Models and Technical Blueprint",
        "Guide review pack | Tentative implementation plan, subject to data access",
    )
    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "This document describes the likely technical route for the PhD after data access is confirmed. "
        "It is a blueprint for guide discussion, not a final implementation commitment."
    )
    doc.add_heading("Proposed System Idea", level=1)
    doc.add_paragraph(
        "Build an explainable clinical decision-support framework that predicts IVF outcome probability, explains patient-specific factors, "
        "and supports doctor-patient counseling without replacing clinician judgment."
    )
    doc.add_heading("High-Level Workflow", level=1)
    add_table(
        doc,
        ["Step", "Purpose", "Main Output"],
        [
            ["Data access and ethics", "Confirm clinic data, permission, anonymization and outcome availability.", "Approved data access plan."],
            ["Data cleaning", "Handle missing values, duplicates, inconsistent formats and outliers.", "Clean research dataset."],
            ["Variable mapping", "Map demographic, clinical, treatment, embryology and optional lifestyle variables.", "Feature dictionary and analysis dataset."],
            ["Exploratory analysis", "Understand distributions, outcome rates, missingness and class imbalance.", "EDA summary and modelling decisions."],
            ["Model development", "Compare baseline and advanced models.", "Candidate prediction models."],
            ["Explainability", "Explain global patterns and patient-level predictions.", "XAI outputs for clinician review."],
            ["Validation", "Test internal, temporal, subgroup or external performance depending on data.", "Validation report."],
            ["CDSS prototype", "Convert model and explanation into a doctor-facing output.", "Prototype or mock decision-support screen."],
        ],
        [1.45, 2.75, 2.1],
    )
    doc.add_heading("Data Model", level=1)
    add_table(
        doc,
        ["Table / Entity", "Example Fields", "Use"],
        [
            ["Patient", "anonymous patient ID, age, BMI, infertility duration, diagnosis.", "Patient-level risk and counselling context."],
            ["Cycle", "cycle ID, year, IVF/ICSI, fresh/frozen, protocol, trigger, endometrial thickness.", "Treatment-cycle modelling and temporal validation."],
            ["Hormone / ovarian reserve", "AMH, AFC, FSH, LH, estradiol, progesterone.", "Clinical prediction and ovarian response explanation."],
            ["Embryology", "oocytes, MII oocytes, fertilization, embryo grade, blastocyst grade, embryos transferred.", "Multimodal prediction if available."],
            ["Outcome", "clinical pregnancy, live birth, miscarriage, implantation, oocyte yield.", "Dependent variable definition."],
            ["Lifestyle", "smoking, sleep, stress, activity, diet.", "Optional personalization only if reliably collected."],
        ],
        [1.45, 3.0, 1.85],
    )
    doc.add_heading("Candidate Models", level=1)
    add_table(
        doc,
        ["Model", "Role", "When To Use"],
        [
            ["Logistic Regression", "Baseline and interpretable comparison model.", "Always useful as a transparent starting point."],
            ["Decision Tree", "Simple rule-like model.", "Useful for explanation but may overfit."],
            ["Random Forest", "Non-linear tabular-data baseline.", "Useful when relationships are not purely linear."],
            ["XGBoost / LightGBM", "Strong candidate models for structured IVF data.", "Likely primary advanced models for tabular clinical/embryology data."],
            ["SVM", "Comparison model.", "Use only if dataset size and feature scaling support it."],
            ["Neural network", "Deep learning candidate.", "Use only for large data, image/time-lapse data or clear technical justification."],
        ],
        [1.45, 2.25, 2.65],
    )
    doc.add_heading("Recommended Model Strategy", level=1)
    add_numbered(
        doc,
        [
            "Start with Logistic Regression as the baseline.",
            "Add Random Forest as a non-linear comparison model.",
            "Use XGBoost or LightGBM as the main advanced tabular-data model.",
            "Use deep learning only if image, time-lapse or sufficiently large data are available.",
            "Select the final model using performance, calibration, explainability and clinical usefulness, not accuracy alone.",
        ],
    )
    doc.add_heading("Explainability Methods", level=1)
    add_table(
        doc,
        ["Method", "Use", "Caution"],
        [
            ["SHAP", "Global feature importance and patient-level positive/negative drivers.", "Explanations must be interpreted clinically, not as causality."],
            ["LIME", "Optional local explanation comparison.", "Can be unstable; use as supporting method only."],
            ["Permutation importance", "Model-agnostic feature importance.", "Can be affected by correlated predictors."],
            ["Partial dependence / related plots", "Shows model relationship for selected variables.", "Use only for clinically meaningful variables."],
            ["Explanation card", "Doctor-facing summary of probability, key factors and caution notes.", "Must be reviewed by clinicians before claiming usefulness."],
        ],
        [1.35, 2.65, 2.3],
    )
    doc.add_page_break()
    doc.add_heading("Evaluation Metrics", level=1)
    add_table(
        doc,
        ["Area", "Candidate Measures", "Why It Matters"],
        [
            ["Discrimination", "AUC, sensitivity, specificity, precision, recall, F1.", "Shows whether model separates outcome groups."],
            ["Calibration", "Calibration curve and Brier score.", "Important because clinicians need reliable probabilities."],
            ["Clinical usefulness", "Decision curve analysis if feasible.", "Shows whether predictions may support decisions better than simple strategies."],
            ["Generalization", "Internal, temporal, subgroup or external validation.", "Addresses repeated literature gap around poor validation."],
            ["Explainability/usefulness", "SHAP review, clinician feedback or structured questionnaire.", "Tests whether explanations are understandable and useful."],
        ],
        [1.45, 2.35, 2.5],
    )
    doc.add_heading("Possible Technology Stack", level=1)
    add_table(
        doc,
        ["Area", "Candidate Tools", "Comment"],
        [
            ["Data processing", "Python, Pandas, NumPy.", "Suitable for structured IVF data cleaning."],
            ["Machine learning", "Scikit-learn, XGBoost, LightGBM.", "Practical and widely used for tabular clinical data."],
            ["Explainable AI", "SHAP, LIME, permutation importance.", "Use with careful clinical interpretation."],
            ["Visualization", "Matplotlib, Seaborn, Plotly.", "Useful for EDA, model results and guide/panel presentation."],
            ["Prototype", "Streamlit, Flask or FastAPI.", "Prototype only; not clinical deployment."],
            ["Database", "CSV/Excel initially; PostgreSQL/MySQL if a formal app is needed.", "Depends on data source and system scope."],
        ],
        [1.45, 2.3, 2.55],
    )
    doc.add_heading("Expected Decision-Support Output", level=1)
    add_table(
        doc,
        ["Output", "Example", "Purpose"],
        [
            ["Predicted probability", "Estimated probability of clinical pregnancy or live birth.", "Gives model output in understandable form."],
            ["Positive factors", "Age range, AMH, embryo grade or endometrial thickness if supportive.", "Shows factors increasing predicted chance."],
            ["Negative factors", "High BMI, low AMH, thin endometrium etc. if model-supported.", "Shows factors reducing predicted chance."],
            ["Modifiable / non-modifiable split", "BMI or smoking versus age or infertility duration.", "Helps counselling without overclaiming."],
            ["Caution note", "Decision support only; clinician judgment remains final.", "Prevents inappropriate use."],
        ],
        [1.6, 2.65, 2.05],
    )
    doc.add_heading("Dataset-Dependent Technical Choices", level=1)
    add_table(
        doc,
        ["Dataset Situation", "Technical Direction"],
        [
            ["Clinical data only", "Tabular ML with XAI and internal/temporal validation."],
            ["Clinical + embryology", "Multimodal tabular modelling with clinical and embryology feature groups."],
            ["Image/time-lapse data", "Deep learning or hybrid model only if data volume, labels and ethics support it."],
            ["Lifestyle data", "Add questionnaire/context variables; analyze carefully and avoid causal claims."],
            ["Multiple clinics/sources", "External validation and clinic-wise generalization analysis."],
            ["Clinician review possible", "Add structured usability or usefulness evaluation of explanation outputs."],
        ],
        [2.0, 4.3],
    )
    doc.add_heading("Implementation Roadmap", level=1)
    add_table(
        doc,
        ["Work Package", "Output"],
        [
            ["Data access and ethics", "Permission, variable list and anonymization plan."],
            ["Data preparation", "Cleaned dataset and feature dictionary."],
            ["EDA", "Missingness, distributions, outcome rates and imbalance summary."],
            ["Model development", "Baseline and advanced model comparison."],
            ["XAI integration", "Global and patient-level explanation outputs."],
            ["Validation", "Internal/temporal/subgroup/external validation report as feasible."],
            ["Prototype", "Doctor-facing decision-support screen or report."],
            ["Clinician review", "Feedback on explanation clarity and usefulness if feasible."],
        ],
        [2.25, 4.05],
    )
    doc.add_heading("Important Boundary", level=1)
    add_bullets(
        doc,
        [
            "Do not claim the model improves IVF success unless a clinical outcome study proves it.",
            "Do not claim deployment; use prototype or decision-support framework language.",
            "Do not claim causality from retrospective prediction data.",
            "Do not use deep learning as the main plan unless the dataset justifies it.",
            "The defensible contribution is prediction, explanation, validation and clinician-facing decision-support design.",
        ],
    )
    return save_doc(doc, "05_methods_models_and_technical_blueprint.docx")


def rq_hypothesis_stats():
    doc = setup_doc(
        "Research Questions, Hypotheses and Statistical Plan",
        "Guide review pack | Draft academic structure for guide discussion",
    )
    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "This document converts the research direction into testable questions and dataset-dependent hypotheses. "
        "The final RQs and hypotheses should be fixed only after outcome, variables and validation data are confirmed."
    )
    doc.add_heading("Research Question Selection Logic", level=1)
    add_table(
        doc,
        ["Step", "Decision"],
        [
            ["1", "Confirm available outcome: live birth, clinical pregnancy, miscarriage or another IVF target."],
            ["2", "Confirm available predictors: clinical, embryology, lifestyle, imaging or mixed data."],
            ["3", "Confirm validation option: internal, temporal, external, multi-clinic or subgroup validation."],
            ["4", "Confirm whether clinician review of explanation outputs is possible."],
            ["5", "Select only RQs and hypotheses that can be tested with the available data."],
        ],
        [0.55, 5.75],
    )
    doc.add_heading("Core Research Questions", level=1)
    add_table(
        doc,
        ["RQ", "Research Question", "Required Data"],
        [
            ["RQ1", "Which available clinical variables are associated with IVF outcome prediction?", "Clinical variables plus reliable outcome."],
            ["RQ2", "Can ML models predict IVF outcome better than a baseline statistical model?", "Enough records, outcome balance and clinical predictors."],
            ["RQ3", "Does adding embryology variables improve prediction compared with clinical variables alone?", "Clinical plus structured embryology variables."],
            ["RQ4", "Do lifestyle/contextual variables add predictive or personalization value?", "Reliable lifestyle/contextual variables."],
            ["RQ5", "Which variables most influence individual predictions according to XAI methods?", "Trained model plus interpretable features."],
            ["RQ6", "Are model probabilities well calibrated for counseling?", "Predicted probabilities plus observed outcomes."],
            ["RQ7", "Does performance differ across age, diagnosis, transfer type or other subgroups?", "Subgroup variables with sufficient sample size."],
            ["RQ8", "Does the model generalize to a second clinic, time period or source?", "External source or temporal split."],
            ["RQ9", "Are clinician-facing explanations understandable and useful?", "Clinician review protocol plus explanation outputs."],
        ],
        [0.55, 3.55, 2.2],
    )
    doc.add_heading("Research Questions By Dataset Scenario", level=1)
    add_table(
        doc,
        ["Dataset Scenario", "Suitable RQs", "Avoid"],
        [
            ["Clinical data only", "RQ1, RQ2, RQ5, RQ6, selected subgroup RQs.", "Embryology, lifestyle or external validation claims."],
            ["Clinical + embryology", "RQ1, RQ2, RQ3, RQ5, RQ6, RQ7.", "Lifestyle claims unless lifestyle data exists."],
            ["Clinical + lifestyle", "RQ1, RQ2, RQ4, RQ5, RQ6.", "Embryology claims unless embryo data exists."],
            ["Independent validation data", "RQ8 plus relevant model RQs.", "Broad generalization beyond available sources."],
            ["Clinician review possible", "RQ9 plus XAI usefulness questions.", "Claims of clinician usefulness without review."],
            ["Live birth unavailable", "Clinical pregnancy version of RQs.", "Live-birth prediction claims."],
        ],
        [1.65, 2.35, 2.3],
    )
    doc.add_heading("Tentative Hypotheses", level=1)
    add_table(
        doc,
        ["H", "Tentative Hypothesis", "Required Data", "Evaluation"],
        [
            ["H1", "Clinical variables can predict the selected IVF outcome better than chance.", "Clinical predictors plus outcome.", "AUC with confidence interval; baseline model."],
            ["H2", "An advanced ML model improves prediction compared with logistic regression.", "Adequate sample size and outcome balance.", "AUC, calibration, Brier score, cross-validation/test-set comparison."],
            ["H3", "Adding embryology variables improves prediction over clinical variables alone.", "Clinical plus embryology variables.", "Model comparison, calibration and net benefit if feasible."],
            ["H4", "Lifestyle variables add predictive or personalization value beyond clinical variables.", "Reliable lifestyle data.", "Model comparison and feature contribution; avoid causality."],
            ["H5", "XAI identifies clinically meaningful patient-level prediction drivers.", "Trained model plus XAI outputs.", "SHAP/LIME review and clinician interpretation if available."],
            ["H6", "Model performance varies across important patient subgroups.", "Subgroup variables and sufficient sample.", "Stratified AUC/calibration and interaction checks."],
            ["H7", "The model maintains acceptable performance on independent/temporal validation data.", "Second source or later time-period data.", "External/temporal discrimination and calibration."],
            ["H8", "Clinician-facing explanations are perceived as useful for counseling.", "Clinician review study.", "Likert ratings, structured feedback or short interview."],
        ],
        [0.45, 2.6, 1.7, 1.55],
    )
    doc.add_heading("Hypotheses To Avoid For Now", level=1)
    add_table(
        doc,
        ["Avoid Hypothesis", "Why Unsafe", "Safer Version"],
        [
            ["The model will increase IVF success rate.", "Requires prospective clinical intervention and outcome testing.", "The model may support counseling; clinical impact needs future prospective validation."],
            ["Lifestyle causes IVF failure.", "Observational prediction data cannot prove causality.", "Lifestyle variables may be associated with outcomes if measured reliably."],
            ["XAI will make doctors trust the model.", "Trust must be measured, not assumed.", "Clinicians can rate explanation usefulness and understandability."],
            ["The model works for all Indian IVF clinics.", "Requires broad multi-center Indian validation.", "The model can be tested on available Indian data and limitations stated."],
            ["Deep learning will be superior.", "Depends on data type and sample size.", "Deep learning may be considered only if data justifies it."],
        ],
        [1.8, 2.15, 2.35],
    )
    doc.add_heading("Objective To RQ Mapping", level=1)
    add_table(
        doc,
        ["Objective", "Related RQs", "Data Requirement"],
        [
            ["Identify useful IVF outcome predictors.", "RQ1, RQ5", "Clinical variables and outcome."],
            ["Compare baseline and ML models.", "RQ2, RQ6", "Model-development dataset."],
            ["Test value of embryology data.", "RQ3", "Structured embryology variables."],
            ["Test value of lifestyle/contextual data.", "RQ4", "Reliable lifestyle data."],
            ["Evaluate validation/generalization.", "RQ7, RQ8", "Subgroup, temporal or external validation data."],
            ["Evaluate clinician-facing explanations.", "RQ9", "Doctor review protocol."],
        ],
        [2.1, 1.15, 3.05],
    )
    doc.add_heading("Statistical and ML Evaluation Plan", level=1)
    add_table(
        doc,
        ["Question Type", "Suitable Analysis", "Important Note"],
        [
            ["Dataset description", "Descriptive statistics, missingness summary, outcome rate.", "Required before modelling."],
            ["Group comparison", "Chi-square, t-test or Mann-Whitney U as appropriate.", "Depends on variable type and distribution."],
            ["Binary outcome association", "Logistic regression.", "Useful baseline and interpretable model."],
            ["Model discrimination", "AUC/ROC, PR-AUC, sensitivity, specificity.", "AUC alone is not enough."],
            ["Probability reliability", "Calibration curve, Brier score.", "Important for counseling."],
            ["Imbalanced outcome", "Precision, recall, F1, PR-AUC.", "Important if positive outcome class is rare."],
            ["Model comparison", "Cross-validation/test-set metrics; DeLong test if appropriate.", "Avoid data leakage."],
            ["Feature contribution", "SHAP, permutation importance.", "Explanation, not causality."],
            ["Subgroup behavior", "Stratified metrics, interaction terms.", "Needs sufficient sample size."],
            ["Clinical usefulness", "Decision curve analysis if thresholds are meaningful.", "Use only if clinically interpretable."],
            ["Clinician review", "Likert-scale survey and qualitative feedback.", "Requires review protocol."],
        ],
        [1.55, 2.4, 2.35],
    )
    doc.add_heading("Safe Current Position", level=1)
    add_bullets(
        doc,
        [
            "The main hypothesis is not finalized because the dependent variable and predictor set are not confirmed.",
            "A safe tentative hypothesis is that an explainable model using available IVF clinical variables can predict the selected outcome better than chance and provide interpretable patient-level explanations.",
            "If embryology data is available, an additional hypothesis can test whether embryology variables improve prediction over clinical variables alone.",
            "Improved IVF success rate should not be claimed unless future prospective clinical validation proves it.",
        ],
    )
    return save_doc(doc, "06_research_questions_hypotheses_and_stats.docx")


def candidate_topics():
    doc = setup_doc(
        "Candidate PhD Topic Directions",
        "Guide review pack | Topic options for discussion, not final title approval",
    )
    doc.add_heading("Current Position", level=1)
    doc.add_paragraph(
        "The title should not be finalized before clinic data feasibility is confirmed. "
        "The current literature supports a strong direction, but the final wording must depend on outcome availability, data modality and validation feasibility."
    )
    doc.add_heading("Why A Generic Title Is Weak", level=1)
    add_table(
        doc,
        ["Weak Title Direction", "Why It Is Weak", "Stronger Direction"],
        [
            [
                "Machine Learning Model for IVF Success Prediction",
                "Many papers already build IVF prediction models; this sounds like a classifier-comparison study.",
                "Explainable, personalized and clinically usable decision support for IVF outcome prediction.",
            ],
            [
                "Deep Learning for IVF Prediction",
                "Deep learning is data-dependent and should not be promised without image/time-lapse or large-scale data.",
                "Use deep learning only as a conditional method if the dataset justifies it.",
            ],
            [
                "AI to Improve IVF Success",
                "Improved success requires prospective clinical outcome evidence.",
                "AI-assisted prediction, explanation and counseling support.",
            ],
        ],
        [1.8, 2.25, 2.25],
    )
    doc.add_heading("Candidate Directions", level=1)
    add_table(
        doc,
        ["Rank", "Candidate Direction", "Evidence Support", "Risk / Dependency"],
        [
            [
                "1",
                "Explainable and personalized clinical decision-support framework for IVF outcome prediction.",
                "Best overall fit with repeated gaps: XAI, CDSS, personalization, validation and clinical usefulness.",
                "Needs clinic outcome data and careful scope control.",
            ],
            [
                "2",
                "Trustworthy AI-based decision support for ovarian stimulation and IVF outcomes.",
                "Strong connection with dose, trigger and protocol decision-support literature.",
                "Requires detailed stimulation and protocol data.",
            ],
            [
                "3",
                "Explainable multimodal IVF prediction using clinical and embryological data.",
                "Strong if structured embryology variables are available; supports multimodal theme.",
                "Cannot be used if embryology data are unavailable or incomplete.",
            ],
            [
                "4",
                "Integration of clinical, embryological and lifestyle data for explainable IVF success prediction.",
                "Strong personalization angle if lifestyle data can be collected reliably.",
                "Lifestyle data availability and quality are uncertain.",
            ],
            [
                "5",
                "Explainable ML for IVF live-birth prediction with Indian population validation.",
                "Strong Indian-context contribution if data is available.",
                "Depends heavily on access to Indian clinic data and live-birth outcome.",
            ],
        ],
        [0.55, 2.25, 2.05, 1.45],
    )
    doc.add_page_break()
    doc.add_heading("Decision Matrix", level=1)
    add_table(
        doc,
        ["Decision Factor", "Best Current Choice", "Reason"],
        [
            ["Academic novelty", "XAI + CDSS + validation-aware IVF prediction.", "Stronger than only comparing ML algorithms."],
            ["Department fit", "Computer Applications with healthcare informatics and predictive analytics.", "Allows technical contribution and applied clinical relevance."],
            ["Data feasibility", "Keep title flexible until clinic confirms variables.", "Avoids untestable claims."],
            ["Safest outcome wording", "IVF outcome prediction.", "Can later become clinical pregnancy or live birth if data allows."],
            ["Safest data wording", "Clinical and embryological data.", "Use only if embryology variables are confirmed; otherwise clinical data only."],
            ["Safest contribution wording", "Decision-support framework.", "Avoids claiming autonomous decision-making or clinical deployment."],
        ],
        [1.7, 2.25, 2.35],
    )
    doc.add_heading("Safest Working Title", level=1)
    doc.add_paragraph(
        "An Explainable and Personalized Clinical Decision Support Framework for IVF Outcome Prediction Using Multimodal Clinical and Embryological Data"
    )
    doc.add_heading("Fallback Titles By Dataset", level=1)
    add_table(
        doc,
        ["Dataset Confirmed", "Safer Title Direction"],
        [
            ["Clinical data only", "An Explainable Clinical Decision Support Framework for IVF Outcome Prediction Using Clinical Data."],
            ["Clinical + embryology", "An Explainable and Personalized Clinical Decision Support Framework for IVF Outcome Prediction Using Multimodal Clinical and Embryological Data."],
            ["Clinical + embryology + lifestyle", "An Explainable Multimodal Clinical Decision Support Framework for Personalized IVF Outcome Prediction Using Clinical, Embryological and Lifestyle Data."],
            ["Live birth outcome confirmed", "Use live-birth prediction wording if follow-up and sample size are adequate."],
            ["Only clinical pregnancy confirmed", "Use clinical pregnancy prediction wording, not live-birth prediction."],
        ],
        [2.05, 4.25],
    )
    doc.add_heading("Conditional Title Extensions", level=1)
    add_table(
        doc,
        ["Extension", "Use Only If", "Avoid If"],
        [
            ["in Indian Women / Indian IVF Settings", "Indian clinic data and validation are available.", "Only literature support exists but no data access."],
            ["Using Clinical, Embryological and Lifestyle Data", "Lifestyle variables are collected reliably and ethically.", "Lifestyle data are absent or weak self-report."],
            ["with External Validation", "Independent clinic/source/time-period data is available.", "Only one-center internal split is available."],
            ["for Live-Birth Prediction", "Live-birth outcome is available with adequate follow-up.", "Only clinical pregnancy is available."],
            ["Using Deep Learning", "Image, time-lapse or large-scale data justify deep learning.", "Only small structured tabular data are available."],
        ],
        [1.9, 2.25, 2.15],
    )
    doc.add_heading("Recommended Position To Tell Guide", level=1)
    doc.add_paragraph(
        "The recommended position is to present one strongest working title and keep dataset-sensitive title words conditional. "
        "The strongest direction is explainable and personalized clinical decision support for IVF outcome prediction. "
        "The exact outcome, data modality and validation wording should be finalized after clinic discussion."
    )
    doc.add_heading("What We Need From Guide", level=1)
    add_table(
        doc,
        ["Guide Decision", "Why It Matters"],
        [
            ["Should the topic focus on outcome prediction or ovarian stimulation decision support?", "This decides the required clinical variables and model target."],
            ["Should the proposal title stay broad initially?", "This reduces risk before data access is confirmed."],
            ["Should Indian validation be title wording or objective/future scope?", "This depends on likely Indian dataset access."],
            ["Is CDSS/framework wording acceptable for the department?", "It positions the work beyond classifier comparison."],
            ["Which words should be avoided in the first formal proposal?", "Prevents overclaiming before feasibility is known."],
        ],
        [2.5, 3.8],
    )
    doc.add_heading("Guide Input Needed", level=1)
    add_numbered(
        doc,
        [
            "Which candidate direction is academically strongest for the department?",
            "Which direction is practical given likely clinic access?",
            "Should the first proposal emphasize IVF outcome prediction, stimulation decision support or broader CDSS?",
            "Which title words should be avoided until data access is confirmed?",
        ],
    )
    return save_doc(doc, "07_candidate_phd_topic_directions.docx")


def readme():
    text = """# Guide Review Pack

This folder contains compact material to share with the PhD guide for review.

## Files

1. `01_literature_review_matrix.xlsx` - detailed Excel paper matrix.
2. `02_compact_literature_review_summary.docx` - short review summary.
3. `03_theme_and_gap_analysis.docx` - theme classification and repeated gaps.
4. `04_dataset_feasibility_and_data_requirements.docx` - required variables and clinic questions.
5. `05_methods_models_and_technical_blueprint.docx` - tentative technical plan.
6. `06_research_questions_hypotheses_and_stats.docx` - draft RQs, hypotheses and evaluation plan.
7. `07_candidate_phd_topic_directions.docx` - candidate titles and decision dependencies.

## Suggested Sharing Order

Start with the Excel workbook, compact literature summary, theme-gap document, dataset feasibility document and candidate topic directions.

Use the methods/statistics document as supporting material if the guide asks about implementation feasibility.

## Important Position

The final PhD title is intentionally not fixed yet. The safest current direction is explainable and personalized clinical decision support for IVF outcome prediction. Final wording depends on clinic data feasibility.
"""
    path = OUT / "README.md"
    path.write_text(text, encoding="utf-8")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    paths = [
        literature_summary(),
        theme_gap_analysis(),
        dataset_feasibility(),
        methodology_blueprint(),
        rq_hypothesis_stats(),
        candidate_topics(),
    ]
    readme()
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
